import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Cm

def clean_path(p):
    return p.strip().replace('\u202a', '').replace('\u202c', '').strip('"')

def browse_docx():
    path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if path:
        entry_docx.delete(0, tk.END)
        entry_docx.insert(0, clean_path(path))

def browse_folder():
    path = filedialog.askdirectory()
    if path:
        entry_folder.delete(0, tk.END)
        entry_folder.insert(0, clean_path(path))

def browse_output():
    path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if path:
        entry_output.delete(0, tk.END)
        entry_output.insert(0, clean_path(path))

def preview_info():
    try:
        input_docx_path = clean_path(entry_docx.get())
        image_folder = clean_path(entry_folder.get())
        target_column = int(entry_column.get()) - 1

        if not os.path.exists(input_docx_path) or not os.path.isdir(image_folder):
            raise FileNotFoundError

        doc = Document(input_docx_path)
        empty_cells = 0
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) > target_column and not row.cells[target_column].text.strip():
                    empty_cells += 1

        images = sorted([
            f for f in os.listdir(image_folder)
            if f.lower().endswith(('.png', '.jpg', '.jpeg'))
        ])
        total_images = len(images)

        label_preview.config(
            text=f"🧮 Empty cells: {empty_cells}, Images available: {total_images}"
        )

    except Exception as e:
        messagebox.showerror("Error", f"Preview failed: {e}")

def insert_images():
    try:
        input_docx_path = clean_path(entry_docx.get())
        image_folder = clean_path(entry_folder.get())
        output_docx_path = clean_path(entry_output.get())
        target_column = int(entry_column.get()) - 1
        image_width_cm = float(entry_width.get())

        if not os.path.exists(input_docx_path):
            raise FileNotFoundError("Input Word file not found.")
        if not os.path.isdir(image_folder):
            raise FileNotFoundError("Image folder not found.")

        doc = Document(input_docx_path)

        images = sorted([
            os.path.join(image_folder, f)
            for f in os.listdir(image_folder)
            if f.lower().endswith(('.png', '.jpg', '.jpeg'))
        ])
        image_index = 0
        inserted = 0

        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) > target_column:
                    cell = row.cells[target_column]
                    if not cell.text.strip() and image_index < len(images):
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run()
                        run.add_picture(images[image_index], width=Cm(image_width_cm))
                        image_index += 1
                        inserted += 1

        doc.save(output_docx_path)
        messagebox.showinfo("Success", f"✅ {inserted} image(s) inserted!\nSaved to: {output_docx_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Something went wrong:\n{e}")

# GUI Setup
root = tk.Tk()
root.title("Word Image Inserter")

tk.Label(root, text="Input Word file:").grid(row=0, column=0, sticky="e")
entry_docx = tk.Entry(root, width=50)
entry_docx.grid(row=0, column=1)
tk.Button(root, text="Browse", command=browse_docx).grid(row=0, column=2)

tk.Label(root, text="Image folder:").grid(row=1, column=0, sticky="e")
entry_folder = tk.Entry(root, width=50)
entry_folder.grid(row=1, column=1)
tk.Button(root, text="Browse", command=browse_folder).grid(row=1, column=2)

tk.Label(root, text="Output file:").grid(row=2, column=0, sticky="e")
entry_output = tk.Entry(root, width=50)
entry_output.grid(row=2, column=1)
tk.Button(root, text="Save As", command=browse_output).grid(row=2, column=2)

tk.Label(root, text="Column to insert (1-based):").grid(row=3, column=0, sticky="e")
entry_column = tk.Entry(root, width=10)
entry_column.insert(0, "3")
entry_column.grid(row=3, column=1, sticky="w")

tk.Label(root, text="Image width (cm):").grid(row=4, column=0, sticky="e")
entry_width = tk.Entry(root, width=10)
entry_width.insert(0, "6.0")
entry_width.grid(row=4, column=1, sticky="w")

tk.Button(root, text="🔍 Preview", command=preview_info).grid(row=5, column=0, pady=10)
tk.Button(root, text="📎 Insert Images", command=insert_images).grid(row=5, column=1, pady=10)

label_preview = tk.Label(root, text="🧮 Preview: Fill in paths and click Preview")
label_preview.grid(row=6, column=0, columnspan=3, pady=5)

root.mainloop()
